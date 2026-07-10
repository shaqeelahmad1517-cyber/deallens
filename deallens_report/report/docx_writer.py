"""Optional Word (.docx) writer using python-docx.

This is the one side-effecting helper (it writes a file) and the one optional
dependency. It's isolated here so the core renderers stay pure and dependency-free.
Install with:  pip install python-docx
"""
from __future__ import annotations

from typing import Any, Dict, Optional

from .normalize import extract

ACCENT = "1F4E79"
ACCENT2 = "2E75B6"


def available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _money(n):
    try:
        return f"${float(n):,.0f}"
    except (TypeError, ValueError):
        return "—" if n is None else str(n)


def write_docx(result: Dict[str, Any], path: str, options: Optional[Dict[str, Any]] = None) -> str:
    """Write a Word report to ``path``. Requires python-docx. Returns the path."""
    try:
        import docx
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required for .docx output: pip install python-docx") from exc

    o = options or {}
    d = extract(result)
    rr = d["recommended_range"]
    norm = d["normalization"]
    ap = d["approaches"]

    doc = docx.Document()

    brand = doc.add_paragraph()
    run = brand.add_run("DEALLENS")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    h = doc.add_heading(f"Valuation Report — {d['target_name']}", level=0)
    for r in h.runs:
        r.font.color.rgb = RGBColor.from_string(ACCENT)

    meta_bits = []
    if o.get("as_of"):
        meta_bits.append(f"As of {o['as_of']}")
    if o.get("prepared_by"):
        meta_bits.append(f"Prepared by {o['prepared_by']}")
    if meta_bits:
        p = doc.add_paragraph("  ·  ".join(meta_bits))
        p.runs[0].italic = True

    doc.add_heading("Recommended Value Range", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"{_money(rr.get('low'))} – {_money(rr.get('high'))}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(ACCENT2)
    doc.add_paragraph(f"Midpoint {_money(rr.get('mid'))}")

    # Summary table
    summary = [
        ("Normalized EBITDA", _money(norm.get("normalized_ebitda"))),
        ("Seller's Discretionary Earnings (SDE)", _money(norm.get("sde"))),
        ("Risk multiple discount", f"{(d['risk'].get('multiple_discount') or 0)*100:.1f}%"),
    ]
    if d["diligence"].get("completion_pct") is not None:
        summary.append(("Diligence completion", f"{d['diligence']['completion_pct']}%"))
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in summary:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v

    # Approaches
    doc.add_heading("Valuation Approaches", level=1)
    at = doc.add_table(rows=1, cols=2)
    at.style = "Light Grid Accent 1"
    at.rows[0].cells[0].text = "Approach"
    at.rows[0].cells[1].text = "Result"
    if "income" in ap:
        inc = ap["income"]
        for label, key in (("Income — DCF", "dcf"), ("Income — Capitalization", "capitalization")):
            c = at.add_row().cells
            c[0].text = label
            c[1].text = _money(inc.get(key, {}).get("value"))
    if "market" in ap:
        mk = ap["market"]
        c = at.add_row().cells
        c[0].text = f"Market — {str(mk.get('metric','')).upper()} {mk.get('low_multiple')}–{mk.get('high_multiple')}×"
        c[1].text = f"{_money(mk.get('low'))} – {_money(mk.get('high'))}"
    if "asset" in ap:
        c = at.add_row().cells
        c[0].text = "Asset — Net Asset Value"
        c[1].text = _money(ap["asset"].get("value"))

    edr = d.get("effective_discount_rate")
    if edr:
        basis = "public-company" if edr <= 0.12 else "small-business"
        note = doc.add_paragraph(
            f"Income methods use an effective discount rate of {edr*100:.1f}% — {basis} cost of capital.")
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(9)

    # Comparables
    if d["comparables"]:
        c = d["comparables"]
        m = c.get("modifiers", {})
        doc.add_heading("Comparables Basis", level=1)
        doc.add_paragraph(
            f"Sector {c.get('sector_matched')}, {str(c.get('metric','')).upper()} base "
            f"{c.get('base_band')} → applied {c.get('low_multiple')}–{c.get('high_multiple')}× "
            f"(size ×{m.get('size_factor')}, growth ×{m.get('growth_factor')})."
        )

    # --- Comprehensive Due Diligence -------------------------------------
    dl = d.get("diligence", {}) or {}
    ai_findings = dl.get("ai_findings") or []
    risk_profile = [p for p in (dl.get("risk_profile") or []) if p.get("level") not in (None, "none")]
    red_flags = d.get("red_flags") or []
    if ai_findings or red_flags or risk_profile:
        doc.add_heading("Due Diligence", level=1)
        comp = dl.get("completion_pct")
        if comp is not None and comp <= 0:
            note = doc.add_paragraph(
                "Checklist not yet worked through — the items below are auto-detected from "
                "the financials and the uploaded document, and should be verified before "
                "relying on this valuation.")
            note.runs[0].italic = True

        if ai_findings:
            doc.add_heading("Findings from the document", level=2)
            doc.add_paragraph("Auto-read from the uploaded report — these provisionally reduce the "
                              "value; confirm each on the checklist for full weight.").runs[0].italic = True
            for f in ai_findings:
                doc.add_paragraph(
                    f"[{str(f.get('severity','')).upper()}] {f.get('category','')}: "
                    f"{f.get('finding') or f.get('label','')}", style="List Bullet")

        if red_flags:
            doc.add_heading("Red flags", level=2)
            doc.add_paragraph("Triggered by the deal's key facts.").runs[0].italic = True
            for f in red_flags:
                doc.add_paragraph(
                    f"[{str(f.get('severity','')).upper()}] {f.get('category','')}: {f.get('label','')}",
                    style="List Bullet")

        if risk_profile:
            order = {"high": 3, "medium": 2, "low": 1}
            doc.add_heading("Risk by area", level=2)
            rt = doc.add_table(rows=1, cols=3)
            rt.style = "Light Grid Accent 1"
            for i, h in enumerate(("Area", "Concern", "Open items")):
                rt.rows[0].cells[i].text = h
            for p in sorted(risk_profile, key=lambda p: -order.get(p.get("level"), 0)):
                cells = rt.add_row().cells
                cells[0].text = str(p.get("category", ""))
                cells[1].text = str(p.get("level", "")).title()
                cells[2].text = str(p.get("open_items", p.get("items", "")))

    # Sensitivity
    sens = d["sensitivity"].get("discount_rate") if d["sensitivity"] else None
    if sens:
        doc.add_heading("Sensitivity — DCF vs. Discount Rate", level=1)
        st = doc.add_table(rows=1, cols=2)
        st.style = "Light Grid Accent 1"
        st.rows[0].cells[0].text = "Discount rate"
        st.rows[0].cells[1].text = "DCF value"
        for r in sens:
            cells = st.add_row().cells
            cells[0].text = f"{(r.get('discount_rate') or 0)*100:.1f}%"
            cells[1].text = _money(r.get("dcf_value"))

    if d["disclaimer"]:
        p = doc.add_paragraph(d["disclaimer"])
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(8)

    doc.save(path)
    return path
